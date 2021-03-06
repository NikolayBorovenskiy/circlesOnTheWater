#!/usr/bin/python
# -*- coding: utf-8 -*-
import wx
import wx.lib.scrolledpanel as scrolled
import os
import sys
import logging

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.text.run import Font, Run
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from threading import Thread
from time import *

from core import start_upwork
from utils.common import *
from utils.SQlite3 import *
from utils.models import *
from utils.constants import MAC_ADDRESS, REG_EMAIL, REG_LOGIN, REG_PASSWORD, REG_TEST_NAME, BASE_DIR


#==========================DOCX=====================================
#Save in the docx file.
#Used for this package python-docx
def saveInFile(data, path):
    #check that the transmitted data - valid
    if not len(data) or data[0]<5:
        return "Pass data not valid!"
    document = Document()
    style = document.styles['IntenseQuote']
    font = style.font
    font.color.rgb = RGBColor(255, 0, 0)
    font.size = Pt(10)
    font.name = 'Arial'

    styles = document.styles
    style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
    style = document.styles['Question']
    fontQuestion = style.font
    fontQuestion.color.rgb = RGBColor(51, 153, 0)
    fontQuestion.size = Pt(13)

    style = document.styles['Heading 1']
    font = style.font
    font.color.rgb = RGBColor(102, 51, 102)
    font.size = Pt(18)
    font.bold = True

    document.add_picture(resource_path(os.path.join('images', 'imgo.jpeg')), width=Inches(1.25))
    #name of the test
    paragraph = document.add_paragraph('{}'.format(data[0][1]), style='Heading 1')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Render all the data
    count=0
    for record in data:
        count+=1
        document.add_paragraph('\n#{}. {}'.format(count, record[2]), style = 'Question')
        if record[5]:

            document.add_paragraph('There may be more than one answer.', style='IntenseQuote')
        counterAnswer = 0
        for answer in record[3].split('#~'):
            counterAnswer+=1
            #Select the correct answer in bold
            #print '----------------------'
            #print answer.strip()
            #print [item.strip() for item in record[4].split('#~')]
            #print '----------------------'

            if answer.strip() in [item.strip() for item in record[4].split('#~')]:
                p = document.add_paragraph("")
                p.add_run("\n{}. {}".format(counterAnswer, answer)).bold = True
            else:
                p = document.add_paragraph("\n{}. {}".format(counterAnswer, answer))

    document.save('{}.docx'.format(path))
    return "File success write."


def readDocx(path):
    document = Document(path)
    paragraphs = document.paragraphs()
    text = [i.text.encode('utf-8') for i in paragraphs]

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile


"""
Module that extract text from MS XML Word document (.docx).
(Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
"""

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'

def get_docx_text(path):
    pass
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [node.text
                 for node in paragraph.getiterator(TEXT)
                 if node.text]
        if texts:
            paragraphs.append(''.join(texts))

    return paragraphs


def parse_docx(data):
    #Connect to database
    path = resource_path(os.path.join('data', 'database.db'))
    cur, con = connect_or_create(path)
    qestionText = ''
    testName = data[0]
    qestionPosition = [i for i in range(len(data)) if data[i].startswith('#')]
    mapDict = {}
    ID = None
    for i in range(len(data)):
        if i in qestionPosition:
            try:
                qestionText = re.sub(r'#\d{0,9}. {0,}', '', data[i])
                ID = filter_table("Qestion", cur, "TEST", "QESTION", [testName, qestionText])[0][0]
            except:
                result = filter_table("Qestion", cur, "TEST", None, [testName])

                for j in result:
                    if qestionText == re.sub(r'\n', '', j[2]):
                        ID = j[0]
            mapDict[ID] = ''
        if data[i].startswith('+'):
            mapDict[ID] += (re.match(r' {0,}\+ {0,}\d+', data[i]).group()[1])

    for i in mapDict.keys():
        if mapDict[i]:
            questionObj = Qestion(*list(filter_table("Qestion", cur, "ID", None, [i])[0])[1:])
            questionObj.findAnswer(mapDict[i])
            try:
                print update_record("Qestion", cur, con, "CORRECT", '#~'.join(questionObj.correctAnswer), i)
            except:
                print "Save record error."

#===================================================================


#===================================================================================================
class Event(object):
    def __init__(self):
        self.Id = None


#===================================================================================================
class ChoseTestDialog(wx.Dialog):
    def __init__(self, parent, id=-1, title=None, tests=None):
        logger.info('Stage 13')
        wx.Dialog.__init__(self, parent, id, title, size=(330, 340))

        self.mainSizer = wx.BoxSizer(wx.VERTICAL)
        self.buttonSizer = wx.BoxSizer(wx.HORIZONTAL)

        #Input field User login
        self.labelInfo = wx.StaticText(self, label="Which test you want to pass?")

        self.allFindedTests = wx.ListBox(self, 26, wx.Point(170,10), wx.Size(290,250), [i.replace('\n', '') for i in tests], wx.LB_SINGLE)
        self.Bind(wx.EVT_LISTBOX, self.OnSelect, id=26)

        self.okButton = wx.Button(self, label="OK", id=wx.ID_OK)
        self.Bind(wx.EVT_BUTTON, self.onOK, id=wx.ID_OK)

        self.mainSizer.Add(self.labelInfo, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.allFindedTests, 0, wx.ALL, 8 )
        self.buttonSizer.Add(self.okButton, 0, wx.ALIGN_CENTER, 20 )
        self.mainSizer.Add(self.buttonSizer, 0, wx.ALIGN_CENTER|wx.ALIGN_BOTTOM, 0)

        self.SetSizer(self.mainSizer)
        self.result = None

    #----------------------------------------------------------------------
    def onOK(self, event):
        try:
            self.result = self.allFindedTests.GetString(self.index).split('.')[0]
        except AttributeError:
            logger.debug("Attribute error")
        if self.result:
            self.Destroy()

    #----------------------------------------------------------------------
    def OnSelect(self, event):
        self.index = event.GetSelection()


#===================================================================================================
class NewUserDialog(wx.Dialog):
    def __init__(self, parent, id=-1, title="Enter Name!"):
        logger.info('Stage 12')
        wx.Dialog.__init__(self, parent, id, title, size=(330, 340))

        self.mainSizer = wx.BoxSizer(wx.VERTICAL)
        self.buttonSizer = wx.BoxSizer(wx.HORIZONTAL)

        #Field of user login
        self.labelLogin = wx.StaticText(self, label="Enter user name:")
        self.fieldLogin = wx.TextCtrl(self, value="", size=(300, 30))

        #Field of user email
        self.labelEmail = wx.StaticText(self, label="Enter email:")
        self.fieldEmail = wx.TextCtrl(self, value="", size=(300, 30))

        #Field of user password
        self.labelPassword = wx.StaticText(self, label="Enter password:")
        self.fieldPassword = wx.TextCtrl(self, value="", size=(300, 30))

        self.okButton = wx.Button(self, label="OK", id=wx.ID_OK)
        self.cancelButton = wx.Button(self, label="CANCEL", id=wx.ID_CANCEL)

        self.mainSizer.Add(self.labelLogin, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.fieldLogin, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.labelEmail, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.fieldEmail, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.labelPassword, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.fieldPassword, 0, wx.ALL, 8 )

        self.buttonSizer.Add(self.okButton, 0, wx.ALL, 20 )
        self.buttonSizer.Add(self.cancelButton, 0, wx.ALL, 20 )

        self.mainSizer.Add(self.buttonSizer, 0, wx.ALIGN_CENTER|wx.ALIGN_BOTTOM, 0)

        self.Bind(wx.EVT_BUTTON, self.onOK, id=wx.ID_OK)
        self.Bind(wx.EVT_BUTTON, self.onCancel, id=wx.ID_NO)

        self.SetSizer(self.mainSizer)
        self.result = None

    #----------------------------------------------------------------------
    def onOK(self, event):
        global _timer1
        _timer1.Start(2000)
        login = self.fieldLogin.GetValue()
        email = self.fieldEmail.GetValue()
        password = self.fieldPassword.GetValue()
        #Validation of entered data
        if validation(login, REG_LOGIN) and validation(email, REG_EMAIL) and validation(password, REG_PASSWORD):
            self.result = []
            self.result.append(login)
            self.result.append(email)
            self.result.append(password)
            self.Close()

    #----------------------------------------------------------------------
    def onCancel(self, event):
        self.result = None
        self.Close()


#===================================================================================================
class StartTestDialog(wx.Dialog):
    def __init__(self, parent, id=-1, title="Start test!"):
        logger.info('Stage 11')
        wx.Dialog.__init__(self, parent, id, title, size=(330, 270))
        global _cur, _con
        self.mainSizer = wx.BoxSizer(wx.VERTICAL)
        self.buttonSizer = wx.BoxSizer(wx.HORIZONTAL)

        #Input field name test
        self.labelTestName = wx.StaticText(self, label="Enter test name")
        self.fieldTestName = wx.TextCtrl(self, value="", size=(300, 30))

        #Input field user
        self.labelUser = wx.StaticText(self, label="Chose user")
        self.listUser = wx.ComboBox(self, -1, pos=(50, 170), size=(300, -1), choices=[i[0] for i in show_table("User", _cur, "USERNAME")], style=wx.CB_READONLY)
        self.setSpeed = wx.CheckBox(self, label='Fast', pos=(20, 20))
        self.setSpeed.SetValue(False)
        self.okButton = wx.Button(self, label="GO", id=wx.ID_YES)
        self.cancelButton = wx.Button(self, label="CANCEL", id=wx.ID_CANCEL)

        self.mainSizer.Add(self.labelTestName, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.fieldTestName, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.labelUser, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.listUser, 0, wx.ALL, 8 )
        self.mainSizer.Add(self.setSpeed, 0, wx.ALL, 8 )

        self.buttonSizer.Add(self.okButton, 0, wx.ALL, 20 )
        self.buttonSizer.Add(self.cancelButton, 0, wx.ALL, 20 )

        self.mainSizer.Add(self.buttonSizer, 0, wx.ALIGN_CENTER|wx.ALIGN_BOTTOM, 0)

        self.Bind(wx.EVT_BUTTON, self.pressOK, id=wx.ID_YES)
        self.Bind(wx.EVT_BUTTON, self.onCancel, id=wx.ID_NO)

        self.SetSizer(self.mainSizer)
        self.result = None

    #----------------------------------------------------------------------
    def pressOK(self, event):
        global _timer
        _timer.Start(2000)
        #Control the user from the database
        global _cur, _con
        user = None
        try:
            user = self.listUser.GetValue()
        except:
            pass
        test = self.fieldTestName.GetValue()
        speed = self.setSpeed.GetValue()

        #Creating a new program stream
        #Run the script in a different thread
        if validation(test, REG_TEST_NAME) and user:
            userName, email, password = None, None, None
            try:
                _, userName, email, password = filter_table("User", _cur, "USERNAME", None, [user])[0]
                file = open(resource_path(os.path.join('data', 'botPhrase.txt')), 'w')
                file.write('Hello!')
                file.close()
            except IndexError:
                logger.debug("IndexError")
            if (userName or email) and password:

                t1 = Thread(target=start_upwork, args=(test.replace(' ', '_'),
                                                               userName.replace(' ', '_'),
                                                               email, password.replace(';', '\;'),
                                                               speed))

                t1.start()
                self.Close()

    #----------------------------------------------------------------------
    def onCancel(self, event):
        self.result = None
        self.Close()


#===================================================================================================
class SolvingDialog(wx.Dialog):
    def __init__(self, parent, id=-1, title="Solving...", testName=None):
        logger.info('Stage 10')
        wx.Dialog.__init__(self, parent, id, title, size=(700, 540),
                            style=wx.DEFAULT_DIALOG_STYLE|wx.RESIZE_BORDER)

        self.testName = testName
        self.upperPanel = UpperPanelSolving(self)
        self.buttonPanel = ButtonPanelSolving(self)

        self.sizer = wx.BoxSizer(wx.VERTICAL)
        self.sizer.Add(self.upperPanel, 0, wx.EXPAND | wx.ALL, 20)
        self.sizer.Add(self.buttonPanel, 0, wx.BOTTOM | wx.ALL, 30)
        self.SetSizerAndFit(self.sizer)

#============================================================================================
class UpperPanelSolving(wx.Panel):
    def __init__(self,  *args, **kwargs):
        logger.info('Stage 9')
        wx.Panel.__init__(self, *args, **kwargs)
        self.testName = args[0].testName
        global _cur, _con
        self.counter = 1
        self.questionObj = None

        self.cb_list = []
        self.subs = []
        evt = Event()
        self.Change(evt)

    #----------------------------------------------------------------------
    def Change(self, e=None):
        #global _cur, _con
        #Listening to what the button is pressed - Next or Back, if the questions asked, but hide the panel.
        if e.Id == 30:
            self.counter+=1
        if e.Id == 31:
            self.counter-=1
        if self.counter<1:
           self.GetParent().Close()

        self.sizer = wx.BoxSizer(wx.VERTICAL)
        self.sizerQuestion = wx.BoxSizer(wx.VERTICAL)
        self.sizerMoreOneAnswer = wx.BoxSizer(wx.VERTICAL)
        self.sizerAnswer = wx.FlexGridSizer(wx.VERTICAL)

        for a in self.subs:
            a.Destroy()

        self.cb_list = []
        self.subs = []
        #tableIndexs = [i[0] for i in filter_table("Qestion", _cur, "TEST", None, [self.testName])]

        try:
            self.questionObj = Qestion(*(list(filter_table("Qestion", _cur, "TEST", None, [self.testName])[self.counter-1]))[1:])
        except IndexError:
            self.GetParent().Close()

        questionText = wx.StaticText(self, -1, '', style=wx.TE_MULTILINE)
        questionText.SetLabel("{}. {}\n".format(self.counter, hyphenation(self.questionObj.qestionText)))
        self.subs.append(questionText)
        self.sizerQuestion.Add(questionText, 1, wx.CENTER)

        #The inscription in the form of the answer may be more than one
        moreOneAnswer = wx.StaticText(self, -1, '', style=wx.TE_MULTILINE)
        if self.questionObj.moreOneAnswer:
            moreOneAnswer.SetLabel("more than one answers")

        self.subs.append(moreOneAnswer)
        self.sizerMoreOneAnswer.Add(moreOneAnswer, 1, wx.CENTER)

        self.sizer.Add(self.sizerQuestion, 0, wx.ALIGN_CENTER|wx.ALIGN_BOTTOM, 0)
        self.sizer.Add(self.sizerMoreOneAnswer, 0, wx.ALIGN_LEFT|wx.ALIGN_BOTTOM, 0)

        answersList = self.questionObj.answers.split('#~')

        #determine the number of the correct answer, which is in the database, knowing your reply. This is done for reliability in real
        numberAnswer = []
        for correct in self.questionObj.correctAnswer.split('#~'):
            #numbering from zero
            try:
                if correct!='No answer' and correct:
                    numberAnswer.append([i.strip() for i in self.questionObj.answers.split('#~')].index(correct.strip()))
            except IndexError as ex:
                logger.debug("Error. Detail: {}".format(ex))

        for i in range(len(answersList)):
            cb = wx.CheckBox(self, -1, '')
            if i in numberAnswer:
                cb.SetValue(True)
            else:
                cb.SetValue(False)
            self.subs.append(cb)
            self.sizerAnswer.Add(cb, 1, wx.LEFT)

            self.cb_list.append(cb)
            answerText = wx.StaticText(self, -1, '', style=wx.TE_MULTILINE)
            answerText.SetLabel(hyphenation(answersList[i]))
            self.subs.append(answerText)
            self.sizerAnswer.Add(answerText, 1, wx.LEFT)

        self.sizer.Add(self.sizerAnswer, 1, wx.EXPAND)
        self.SetSizerAndFit(self.sizer)
        self.GetParent().Fit()

    #----------------------------------------------------------------------
    def OnSave(self, event):
        global _cur, _con
        #print "Counter: ", self.counter
        #print "ID: ", filter_table("Qestion", _cur, "TEST", None, [self.testName])[self.counter][0]
        answersStr = ""
        for i, check in enumerate(self.cb_list):
            if check.GetValue():
                answersStr+=str(i+1)

        self.questionObj.findAnswer(answersStr)
        tableIndexs = [i[0] for i in filter_table("Qestion", _cur, "TEST", None, [self.testName])]
        update_record("Qestion", _cur, _con, "CORRECT", '#~'.join(self.questionObj.correctAnswer), tableIndexs[self.counter-1])


#===================================================================================================
class ButtonPanelSolving(wx.Panel):
    def __init__(self, *args, **kwargs):
        logger.info('Stage 8')
        wx.Panel.__init__(self, *args, **kwargs)

        self.saveBtn = wx.Button(self, label="SAVE",  size = (130, 35))
        self.saveBtn.Bind(wx.EVT_BUTTON, self.GetParent().upperPanel.OnSave)

        self.backBtn = wx.Button(self, id=31,label="BACK",  size = (130, 35))
        self.backBtn.Bind(wx.EVT_BUTTON, self.GetParent().upperPanel.Change)

        self.nextBtn = wx.Button(self, id=30, label="NEXT",  size = (130, 35))
        self.nextBtn.Bind(wx.EVT_BUTTON, self.GetParent().upperPanel.Change)

        self.sizer = wx.BoxSizer(wx.HORIZONTAL)
        self.sizer.Add(self.backBtn, 0, wx.ALL, 20)
        self.sizer.Add(self.saveBtn, 0, wx.ALL, 20)
        self.sizer.Add(self.nextBtn, 0, wx.ALL, 20)
        self.SetSizerAndFit(self.sizer)


#===================================================================================================
class StartPanel(wx.Panel):
    """"""
    #----------------------------------------------------------------------
    def __init__(self, parent):
        logger.info('Stage 7')
        """Constructor"""
        wx.Panel.__init__(self, parent=parent)

        profilesBtn = wx.Button(self, -1, "USERS", (20, 40), (175, 50))
        allTestsBtn = wx.Button(self, -1, "TESTS", (20, 100), (175, 50))

        #add picture
        self.bitmap = wx.Bitmap(resource_path(os.path.join('images', 'upwork.png')))
        wx.EVT_PAINT(self, self.OnPaint)

        #The event handler pressing
        profilesBtn.Bind(wx.EVT_BUTTON, parent.onSwitchUserPanel)
        allTestsBtn.Bind(wx.EVT_BUTTON, parent.onSwitchAllTestsPanel)

    #Draw a picture
    #----------------------------------------------------------------------
    def OnPaint(self, event):
        dc = wx.PaintDC(self)
        dc.DrawBitmap(self.bitmap, 200, 20)


#===================================================================================================
class TestPanel(wx.Panel):
    """"""
    #----------------------------------------------------------------------
    def __init__(self, parent):
        logger.info('Stage 6')
        """Constructor"""
        wx.Panel.__init__(self, parent=parent)
        self.fileSize = 0

        #Create a timer
        global _timer
        _timer = wx.Timer(self, 1)

        self.startBtn = wx.Button(self, -1, "START", (15, 40), (130, 35))
        self.startBtn.Bind(wx.EVT_BUTTON, self.OnShowCustomDialog)

        self.backBtn = wx.Button(self, -1, "BACK", (15, 80), (130, 35))
        self.backBtn.Bind(wx.EVT_BUTTON, parent.onSwitchAllTestsPanel)

        #Add panel to display information
        self.logger = wx.TextCtrl(self,5, "",
                wx.Point(160,10), wx.Size(330,330),
                wx.TE_MULTILINE | wx.TE_READONLY)

        self.Bind(wx.EVT_TIMER, self.OnTimer, id=1)

    #----------------------------------------------------------------------
    def OnTimer(self, event):
        if os.path.getsize(resource_path(os.path.join('data', 'botPhrase.txt')))!=self.fileSize:
            file = open(resource_path(os.path.join('data', 'botPhrase.txt')), 'r')
            fileContent = file.readlines()
            self.logger.SetValue(('\n'.join(fileContent[::-1])).replace('\n\n', '\n'))
            self.fileSize = os.path.getsize(resource_path(os.path.join('data', 'botPhrase.txt')))

            #Condition when found few tests and to choose any one
            if fileContent.count("I found a few tests.\n") and not fileContent.count("Test selected.\n"):
                choseTestDia = ChoseTestDialog(self, -1, 'Select test',  fileContent[fileContent.index("I found a few tests.\n")+1:])
                val = choseTestDia.ShowModal()
                logger.debug(choseTestDia.result)
                writeTempFile("Test selected.\n{}".format(choseTestDia.result))
                choseTestDia.Destroy()

    #----------------------------------------------------------------------
    def OnShowCustomDialog(self, event):
        dia = StartTestDialog(self, -1, 'Start test!')
        val = dia.ShowModal()
        dia.Destroy()


#===================================================================================================
class AllTestsPanel(wx.Panel):
    """"""
    #----------------------------------------------------------------------
    def __init__(self, parent):
        logger.info('Stage 5')
        """Constructor"""
        wx.Panel.__init__(self, parent=parent)
        self.currentDirectory = os.getcwd()
        self.index = None
        self.wildcard = "Document source (*.doc)|*.docx|" \
                "All files (*.*)|*.*"

        global _timer2, _cur, _con
        _timer2 = wx.Timer(self, 3)

        self.homeBtn = wx.Button(self, -1, "HOME", (15, 40), (130, 35))
        self.homeBtn.Bind(wx.EVT_BUTTON, parent.onSwitchMainPanels)

        self.passBtn = wx.Button(self, -1, "EXAM", (15, 80), (130, 35))
        self.passBtn.Bind(wx.EVT_BUTTON, parent.onSwitchTestPanel)

        self.solveBtn = wx.Button(self, -1, "SOLVE", (15, 120), (130, 35))
        self.solveBtn.Disable()
        self.solveBtn.Bind(wx.EVT_BUTTON, self.OnShowSolvingDialog)

        self.saveBtn = wx.Button(self, -1, "SAVE", (15, 160), (130, 35))
        self.saveBtn.Disable()
        self.saveBtn.Bind(wx.EVT_BUTTON, self.saveFile)

        self.loadBtn = wx.Button(self, -1, "LOAD", (15, 200), (130, 35))
        self.loadBtn.Bind(wx.EVT_BUTTON, self.loadFile)

        self.deleteBtn = wx.Button(self, -1, "DELETE", (15, 240), (130, 35))
        self.deleteBtn.Disable()
        self.deleteBtn.Bind(wx.EVT_BUTTON, self.DeleteTest)

        self.Bind(wx.EVT_TIMER, self.OnTimer, id=3)

        self.allTest = None
        seen = set()
        self.allTestsListBox = wx.ListBox(self, 26, wx.Point(170,10), wx.Size(290,250), [], wx.LB_SINGLE)
        #Panel for short information for each test
        self.testInfoText = wx.StaticText(self, -1, '', pos=(190,270), size=(200, 130), style=wx.TE_MULTILINE)

        #The event handler pressing
        self.Bind(wx.EVT_LISTBOX, self.OnSelect, id=26)

    #-----------------------------------------------------------------------------------
    def OnShowSolvingDialog(self, event):
        dia = SolvingDialog(self, -1, 'Solving', self.allTestsListBox.GetString(self.index))
        val = dia.ShowModal()
        dia.Destroy()

    #-----------------------------------------------------------------------------------
    def OnSelect(self, event):
        self.index = event.GetSelection()
        global _cur, _con
        userInfo = self.allTestsListBox.GetString(self.index)
        #unlock button
        if userInfo:
            self.solveBtn.Enable()
            self.saveBtn.Enable()
            self.deleteBtn.Enable()

        allQestion = filter_table("Qestion", _cur, "TEST", None, [userInfo])
        unAnswered = [i for i in filter_table("Qestion", _cur, "TEST", None, [userInfo]) if i[4]!='No answer' and i[4]!='']
        self.testInfoText.SetLabel("Total: {}\nAnswered: {}\nUnanswered: {}".format(len(allQestion),
                                                                                    len(unAnswered),
                                                                                    len(allQestion) - len(unAnswered)))

    def OnTimer(self, event):
        seen = set()
        self.allTest = [x[0] for x in show_table("Qestion", _cur, "TEST") if x[0] not in seen and not seen.add(x[0])]
        self.allTestsListBox.SetItems(self.allTest)
        if self.allTestsListBox.Items == self.allTest:
            self.solveBtn.Disable()
            self.saveBtn.Disable()
            self.deleteBtn.Disable()
            global _timer2
            _timer2.Stop()

    #-----------------------------------------------------------------------------------
    def saveFile(self, event):
        """
        Create and show the Save FileDialog
        """
        global _cur, _con
        userInfo = self.allTestsListBox.GetString(self.index)
        dlg = wx.FileDialog(
            self, message="Save file as ...",
            defaultDir=self.currentDirectory,
            defaultFile="", wildcard=self.wildcard, style=wx.SAVE
            )
        if dlg.ShowModal() == wx.ID_OK:
            path = dlg.GetPath()
            try:
                saveInFile(filter_table("Qestion", _cur, "TEST", None, [userInfo]), path)
            except Exception as ex:
                pass
        dlg.Destroy()

    #-----------------------------------------------------------------------------------
    def loadFile(self, event):
        """
        Create and show the Open FileDialog
        """
        #currentPath = os.getcwd()
        dlg = wx.FileDialog(
            self, message="Choose a file",
            defaultDir=self.currentDirectory,
            defaultFile="",
            wildcard=self.wildcard,
            style=wx.OPEN | wx.MULTIPLE | wx.CHANGE_DIR
            )
        if dlg.ShowModal() == wx.ID_OK:
            paths = dlg.GetPaths()
            logger.debug("You chose the following file(s):")
            for path in paths:
                parse_docx(get_docx_text(path))
        dlg.Destroy()

    #-----------------------------------------------------------------------------------
    def DeleteTest(self, event):
        global _cur, _con
        userInfo = self.allTestsListBox.GetString(self.index)
        delete_record("Qestion", _cur, _con, "ID", [i[0] for i in filter_table("Qestion", _cur, "TEST", None, [userInfo])])
        self.testInfoText.SetLabel("")


#===================================================================================================
class UserPanel(wx.Panel):
    """"""
    #----------------------------------------------------------------------
    def __init__(self, parent):
        logger.info('Stage 4')
        """Constructor"""
        wx.Panel.__init__(self, parent=parent)
        self.index = None

        global _timer1, _cur, _con

        _timer1 = wx.Timer(self, 2)
        self.homeBtn = wx.Button(self, -1, "HOME", (15, 40), (130, 35))
        self.homeBtn.Bind(wx.EVT_BUTTON, parent.onSwitchMainPanels)

        self.newBtn = wx.Button(self, -1, "NEW", (15, 80), (130, 35))
        self.newBtn.Bind(wx.EVT_BUTTON, self.OnShowCustomDialog)

        self.deleteBtn = wx.Button(self, -1, "DELETE", (15, 120), (130, 35))
        self.deleteBtn.Disable()
        self.deleteBtn.Bind(wx.EVT_BUTTON, self.DeleteUser)

        self.allUsers = wx.ListBox(self, 26, wx.Point(170,10), wx.Size(290,250), [i[0] for i in show_table("User", _cur, "USERNAME")], wx.LB_SINGLE)
        self.Bind(wx.EVT_LISTBOX, self.OnSelect, id=26)
        self.userInfoText = wx.StaticText(self, -1, '', pos=(190,270), size=(200, 130), style=wx.TE_MULTILINE)

        #event handlers
        self.Bind(wx.EVT_TIMER, self.OnTimer1, id=2)

    #-----------------------------------------------------------------------------------
    def OnTimer1(self, event):
        self.allUsers.Set([i[0] for i in show_table("User", _cur, "USERNAME")])

    #-----------------------------------------------------------------------------------
    def OnShowCustomDialog(self, event):
        global _timer1, _cur, _con
        _timer1.Start(2000)
        dia = NewUserDialog(self, -1, 'Create new user')
        val = dia.ShowModal()
        #Record results into a database.
        if dia.result is not None and len(dia.result)>=3:
            save_records("User", _cur, _con, dia.result)
        dia.Destroy()

    #-----------------------------------------------------------------------------------
    def OnSelect(self, event):
        global _timer1, _cur, _con
        _timer1.Stop()
        self.index = event.GetSelection()
        userInfo = self.allUsers.GetString(self.index)
        if userInfo:
            self.deleteBtn.Enable()
        try:
            _, userName, email, password = filter_table("User", _cur, "USERNAME", None, [userInfo])[0]
            self.userInfoText.SetLabel("Name: {}\nEmail: {}\nPassword: {}".format(userName, email, password))
        except IndexError as ex:
            logger.debug("Error: {}".format(ex))

    #-----------------------------------------------------------------------------------
    def DeleteUser(self, event):
        global _cur, _con, _timer1
        userInfo = self.allUsers.GetString(self.index)
        _timer1.Start(2000)
        delete_record("User", _cur, _con, "ID", [i[0] for i in filter_table("User", _cur, "USERNAME", None, [userInfo])])
        self.userInfoText.SetLabel("")


#===================================================================================================
class MainFrame(wx.Frame):
    #-----------------------------------------------------------------------------------
    def __init__(self, parent, title):
        no_resize = wx.DEFAULT_FRAME_STYLE & ~ (wx.RESIZE_BORDER |
                                                wx.RESIZE_BOX |
                                                wx.MAXIMIZE_BOX)
        logger.info('Stage 3')

        wx.Frame.__init__(self, parent, -1, title, pos=(400, 150), size=(500, 350), style=no_resize)
        initialization()
        self.Center()
        self.colur = wx.Colour(0, 0, 0)

        self.panelStart = StartPanel(self)
        self.panelUsers = UserPanel(self)
        self.panelTest = TestPanel(self)
        self.panelAllTests = AllTestsPanel(self)

        self.panelUsers.Hide()
        self.panelTest.Hide()
        self.panelAllTests.Hide()

        self.sizer = wx.BoxSizer(wx.VERTICAL)
        self.sizer.Add(self.panelStart, 1, wx.EXPAND)
        self.sizer.Add(self.panelUsers, 1, wx.EXPAND)
        self.sizer.Add(self.panelTest, 1, wx.EXPAND)
        self.sizer.Add(self.panelAllTests, 1, wx.EXPAND)
        self.SetSizer(self.sizer)

    #-----------------------------------------------------------------------------------
    def onSwitchMainPanels(self, event):
        global _timer
        global _timer1
        global _timer2
        _timer.Stop()
        _timer1.Stop()
        _timer2.Stop()
        self.panelStart.Show()
        self.panelUsers.Hide()
        self.panelTest.Hide()
        self.panelAllTests.Hide()
        self.Layout()

    #-----------------------------------------------------------------------------------
    def onSwitchUserPanel(self, event):
        self.panelUsers.Show()
        self.panelStart.Hide()
        self.panelTest.Hide()
        self.panelAllTests.Hide()
        self.Layout()

    #-----------------------------------------------------------------------------------
    def onSwitchTestPanel(self, event):
        self.panelTest.Show()
        self.panelUsers.Hide()
        self.panelStart.Hide()
        self.panelAllTests.Hide()
        self.Layout()

    #-----------------------------------------------------------------------------------
    def onSwitchAllTestsPanel(self, event):
        global _timer2
        _timer2.Start(500)
        self.panelAllTests.Show()
        self.panelUsers.Hide()
        self.panelStart.Hide()
        self.panelTest.Hide()
        self.Layout()


#===================================================================================================
class App(wx.App):
    def OnInit(self):
        #logger.info('Stage 2')
        frame = MainFrame(None, "Circles on the water")
        frame.Show(True)
        self.SetTopWindow(frame)
        return True


def initialization():
    #User validation
    #authorization(MAC_ADDRESS)

    try:
        os.mkdir(resource_path(os.path.join('data')))
    except OSError:
        logger.debug("Directory 'data' already create")

    path = resource_path(os.path.join('data', 'database.db'))
    global _cur, _con
    _cur, _con = connect_or_create(path)

    #Create a table User, if it does not exist yet
    try:
        create_table("User", _cur, _con, USERNAME = "TEXT", EMAIL = "TEXT", PASSWORD = "TEXT")
        create_table("Qestion", _cur, _con, TEST="TEXT", QESTION="TEXT", ANSWERS="TEXT", CORRECT="TEXT", MOREONE = "BOOLEAN")
    except:
        logger.debug("Tables already create.")


# Run the program
if __name__ == "__main__":

    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)
    #initialization
    #initialization()
    # create a file handler
    handler = logging.FileHandler(resource_path('main.log'))
    handler.setLevel(logging.INFO)

    # create a logging format
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)

    # add the handlers to the logger
    logger.addHandler(handler)


    _timer, _timer1, _timer2, _cur, _con = None, None, None, None, None
    logger.info('Stage 1')
    app = App()
    app.MainLoop()