# -*- coding: utf-8 -*-
import sqlite3

from docx import Document
from docx.shared import Inches

"""
Инструменты для работы с базой данных sqlite3
Так же есть инструмент для печати результатов в docx формате
"""

#Создать или подключиться к базе данных
def connect_or_create(nameDB):
	con = sqlite3.connect(nameDB)
	#Создание обьекта Курсор для взаимодействия с базой данных. Т.е. формирования запросос к базе
	cur = con.cursor()
	return cur, con


#Создать таблицу
def create_table(tableName, cusor, idField, testField, qestionFiels, answersField, correctField, moreOneField):
	cusor.execute('CREATE TABLE {}({} INT, {} TEXT, {} TEXT, {} TEXT, {} TEXT, {} BOOLEAN)'.format(tableName, 
																		 idField,
																		 testField,
																		 qestionFiels,
																		 answersField,
																		 correctField,
																		 moreOneField
																		 ))

	return "Success"


#Записать данные в таблицу. Одну запись
def save_records(tableName, cursor, connection, dataList):
	#ID определяется автоматически
	tempDataList = dataList[:]
	tempDataList.insert(0,len(cursor.execute('SELECT * FROM {}'.format(tableName)).fetchall()))
	tempDataList[-1] = bool(tempDataList[-1])
	print tempDataList
	print len(tempDataList)
	cursor.executemany('INSERT INTO {} VALUES(?,?,?,?,?,?)'.format(tableName), (tempDataList,))
	connection.commit()
	return "Success"


#Обновить запись в таблице
def update_record(tableName, cusor, connection, field, value, idNumber):
	cusor.execute("""UPDATE {} SET {} = ? WHERE ID= ?""".format(tableName, field),(value, idNumber))
	connection.commit()
	return "Success"



#Показать содержимое таблицы
def show_table(tableName, cursor, field=None):
	#Посмотреть сожержание таблицы.
	if field == None:
		cursor.execute('SELECT * FROM {}'.format(tableName))
	else:
		cursor.execute('SELECT {} FROM {}'.format(field, tableName))
	data = cursor.fetchall()
	return data

#Фильтрация данных в базе
def filter_table(tableName, cursor, field1, field2=None, dataList=None):
	if field2 != None:
		cursor.execute("SELECT * FROM {} WHERE {} = ? and {} = ?".format(tableName, field1, field2), dataList)
	else:
		cursor.execute("SELECT * FROM {} WHERE {} = ?".format(tableName, field1), dataList)
	data = cursor.fetchall()
	return data


#Распасить модель Qestion
def parserModel(obj):
    results = []
    for i in '{}'.format(obj).split(';'):
        i = i.strip()
        #if i == '{}'.format(obj).split(';')[-1].strip():
        #    i = bool(i)
        results.append(i)
    return results


#Сохранить результаты в docx файле.
#Использую для этого пакет python-docx
def saveInFile(data):
	#проверка, что переданные данные - валидные
	if not len(data) or data[0]<5:
		return "Pass data not valid!"
	document = Document()
	document.add_picture('python_doc/games.jpg', width=Inches(1.25))
	document.add_heading('{}'.format(data[0][1]), 0)
	#Отрендерить все данные
	for record in data:
		document.add_heading('\n{}. {}'.format(record[0]+1, record[2]), level=3)
		if record[5]:
			document.add_paragraph('There may be more than one answer.', style='IntenseQuote')
		counterAnswer = 0
		for answer in record[3].split('#~'):
			counterAnswer+=1
			#Выделяем правильный ответ жирны
			if answer in record[4].split('#~'):
				p = document.add_paragraph("")
				p.add_run("\n{}. {}".format(counterAnswer, answer)).bold = True
			else:
				p = document.add_paragraph("\n{}. {}".format(counterAnswer, answer))

	document.save('{}_qestions.docx'.format(data[0][1]))
	return "File success write."

