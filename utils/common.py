#!/usr/bin/python
# -*- coding: utf-8 -*-
import sys
import os
import commands
import random
import re

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.text.run import Font, Run
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.SQlite3 import *
from utils.models import Qestion


def locationControl(text):
    if text == "I am lost. ;(":
        sys.exit()

#Rear computer mac-adress Linux
def getmac(iface):
    words = commands.getoutput("ifconfig " + iface).split() 
    for x in range(0,len(words)):
        if words[x].strip() == "HWaddr":
            mac = words[x+1].strip()
            break
    if len(mac) == 0:
        mac = 'Mac not found'
    mac = mac[:17]
    return mac


#Rear computer mac-adress Windows
def get_macaddress(host='localhost'):
    """ Returns the MAC address of a network host, requires >= WIN2K. """
    # http://aspn.activestate.com/ASPN/Cookbook/Python/Recipe/347812
    import ctypes
    import socket
    import struct
 
    # Check for api availability
    try:
        SendARP = ctypes.windll.Iphlpapi.SendARP
    except:
        raise NotImplementedError('Usage only on Windows 2000 and above')
 
    # Doesn't work with loopbacks, but let's try and help.
    if host == '127.0.0.1' or host.lower() == 'localhost':
        host = socket.gethostname()
 
    # gethostbyname blocks, so use it wisely.
    try:
        inetaddr = ctypes.windll.wsock32.inet_addr(host)
        if inetaddr in (0, -1):
            raise Exception
    except:
        hostip = socket.gethostbyname(host)
        inetaddr = ctypes.windll.wsock32.inet_addr(hostip)
 
    buffer = ctypes.c_buffer(6)
    addlen = ctypes.c_ulong(ctypes.sizeof(buffer))
    if SendARP(inetaddr, 0, ctypes.byref(buffer), ctypes.byref(addlen)) != 0:
        raise WindowsError('Retreival of mac address(%s) - failed' % host)
 
    # Convert binary data into a string.
    macaddr = ''
    for intval in struct.unpack('BBBBBB', buffer):
        if intval > 15:
            replacestr = '0x'
        else:
            replacestr = 'x'
        if macaddr != '':
            macaddr = ':'.join([macaddr, hex(intval).replace(replacestr, '')])
        else:
            macaddr = ''.join([macaddr, hex(intval).replace(replacestr, '')])
 
    return macaddr


#Authorization of the program through a Mac-address.Run computer with certain MAC-address. For Linux Windows
def authorization(data, iface='eth0'):
    print sys.platform
    print getmac(iface)
    mac = None
    if sys.platform == "win32":
        mac = get_macaddress(host='localhost')
    if sys.platform == "linux2":
        mac = getmac(iface)
    if not mac in data:
        print "Unfamiliar MAC address!"
        sys.exit()


#Save in the docx file.
#Used for this package python-docx
def saveInFile(data, path):
    #check that the transmitted data - valid
    if not len(data) or data[0]<5:
        return "Pass data not valid!"
    document = Document()
    style = document.styles['IntenseQuote']
    font = style.font
    font.color.rgb = RGBColor(255, 0, 0)
    font.size = Pt(10)
    font.name = 'Arial'


    styles = document.styles
    style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
    style = document.styles['Question']
    fontQuestion = style.font
    fontQuestion.color.rgb = RGBColor(51, 153, 0)
    fontQuestion.size = Pt(13)

    style = document.styles['Heading 1']
    font = style.font
    font.color.rgb = RGBColor(102, 51, 102)
    font.size = Pt(18)
    font.bold = True
    
    document.add_picture(os.path.join(os.getcwd(), 'images', 'imgo.jpeg'), width=Inches(1.25))
    #name of the test
    paragraph = document.add_paragraph('{}'.format(data[0][1]), style='Heading 1')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #Render all the data
    count=0
    for record in data:
        count+=1
        document.add_paragraph('\n#{}. {}'.format(count, record[2]), style = 'Question')
        if record[5]:

            document.add_paragraph('There may be more than one answer.', style='IntenseQuote')
        counterAnswer = 0
        for answer in record[3].split('#~'):
            counterAnswer+=1
            #Select the correct answer in bold
            if answer.strip() in [item.strip() for item in record[4].split('#~')]:
                p = document.add_paragraph("")
                p.add_run("\n{}. {}".format(counterAnswer, answer)).bold = True
            else:
                p = document.add_paragraph("\n{}. {}".format(counterAnswer, answer))

    document.save('{}.docx'.format(path))
    return "File success write."

#run the script
def execute(command, *args, **kwargs):
    os.system(command.format(*args, **kwargs))


#Line break
def hyphenation(text, length=100):
    newText=''
    while True:
        if len(text)>length and not text.count('\n'):
            findSpace = text[length:].strip().find(' ')
            if findSpace>0:
                newText += '{}\n'.format(text[:length+findSpace+1])
                text = text[length+findSpace+1:].strip()
            else:
                newText+=text
                break
        else:
            newText+=text
            break
    return newText


#Parse model Qestion
def parserModel(obj):
    results = []
    for i in '{}'.format(obj).split('&;'):
        i = i.strip()
        #if i == '{}'.format(obj).split(';')[-1].strip():
        #    i = bool(i)
        results.append(i)
    return results


def readDocx(path):
    document = Document(path)
    paragraphs = document.paragraphs()
    text = [i.text.encode('utf-8') for i in paragraphs]
    print text

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile


"""
Module that extract text from MS XML Word document (.docx).
(Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
"""

WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'

def get_docx_text(path):
    pass
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [node.text
                 for node in paragraph.getiterator(TEXT)
                 if node.text]
        if texts:
            paragraphs.append(''.join(texts))

    return paragraphs


def parse_docx(data, currentPath):
    #Connect to database
    path = os.path.join(currentPath, 'data', 'database.db')
    cur, con = connect_or_create(path)
    qestionText = ''
    testName = data[0]
    qestionPosition = [i for i in range(len(data)) if data[i].startswith('#')]
    mapDict = {}
    ID = None
    for i in range(len(data)):
        if i in qestionPosition:
            try:
                qestionText = re.sub(r'#\d{0,9}. {0,}', '', data[i])
                ID = filter_table("Qestion", cur, "TEST", "QESTION", [testName, qestionText])[0][0]
            except:
                result = filter_table("Qestion", cur, "TEST", None, [testName])
                
                for j in result:
                    if qestionText == re.sub(r'\n', '', j[2]):
                        ID = j[0]
            mapDict[ID] = ''
        if data[i].startswith('+'):
            mapDict[ID] += (re.match(r' {0,}\+ {0,}\d+', data[i]).group()[1])
            
    for i in mapDict.keys():
        if mapDict[i]:
            questionObj = Qestion(*list(filter_table("Qestion", cur, "ID", None, [i])[0])[1:])
            questionObj.findAnswer(mapDict[i])  
            try:
                print update_record("Qestion", cur, con, "CORRECT", '#~'.join(questionObj.correctAnswer), i)
            except:
                print "Save record error."


def validation(target, regx):
    pattern = re.compile(regx)
    result = pattern.match(target)
    if result and result.group()==target:
        return target
    else:
        return None

# Run the program
#if __name__ == "__main__":
    #for i in get_docx_text("../node.docx"):
    #    print i
